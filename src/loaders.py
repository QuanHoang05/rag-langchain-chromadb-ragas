import os
import glob
import re
import unicodedata
from typing import List
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
import docx  # Sử dụng python-docx để đọc file Word (.docx)

def clean_vietnamese_text(text: str) -> str:
    """
    Hàm làm sạch và chuẩn hóa văn bản tiếng Việt.
    1. Chuẩn hóa Unicode về dạng NFC để tránh lỗi hiển thị/tìm kiếm font tiếng Việt.
    2. Loại bỏ các ký tự điều khiển lạ (trừ xuống dòng, tab).
    3. Gộp các khoảng trắng thừa và các dòng trống liên tiếp.
    """
    if not text:
        return ""
        
    # Chuẩn hóa Unicode NFC
    text = unicodedata.normalize('NFC', text)
    
    # Loại bỏ các ký tự điều khiển (Unicode category 'C' - control characters)
    # Giữ lại \n và \t phục vụ cấu trúc câu
    text = "".join(
        char for char in text
        if not unicodedata.category(char).startswith('C') or char in '\n\t'
    )
    
    # Gộp khoảng trắng thừa thành khoảng trắng đơn
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Gộp các dòng trống liên tiếp (3 dòng trở lên thành tối đa 2 dòng trống)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()

class SimpleLoader:
    """
    Lớp tải tài liệu hỗ trợ PDF và Word (.docx) kèm theo dán nhãn metadata (Metadata Tagging).
    """
    
    def __init__(self):
        pass
        
    def load_pdf(self, file_path: str) -> List[Document]:
        """
        Nạp file PDF sử dụng PyPDFLoader và làm sạch nội dung văn bản.
        """
        if not os.path.exists(file_path):
            print(f"[Cảnh báo] File không tồn tại: {file_path}")
            return []
            
        try:
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            
            # Làm sạch page content của từng trang
            for doc in docs:
                doc.page_content = clean_vietnamese_text(doc.page_content)
                # Lưu thông tin tên file vào metadata
                doc.metadata["source_name"] = os.path.basename(file_path)
                
            return docs
        except Exception as e:
            print(f"[Lỗi] Không thể nạp PDF {file_path}: {str(e)}")
            return []
            
    def load_docx(self, file_path: str) -> List[Document]:
        """
        Nạp file Word (.docx) thủ công bằng python-docx để tối ưu hiệu năng và kiểm soát định dạng.
        """
        if not os.path.exists(file_path):
            print(f"[Cảnh báo] File không tồn tại: {file_path}")
            return []
            
        try:
            doc_obj = docx.Document(file_path)
            full_text = []
            
            # Trích xuất văn bản từ các đoạn văn
            for para in doc_obj.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
                    
            # Trích xuất văn bản từ các bảng biểu (nếu có)
            for table in doc_obj.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
                        
            text_content = "\n".join(full_text)
            cleaned_content = clean_vietnamese_text(text_content)
            
            # Tạo đối tượng Document tương đương cấu trúc LangChain
            doc = Document(
                page_content=cleaned_content,
                metadata={
                    "source": file_path,
                    "source_name": os.path.basename(file_path),
                    "page": 1  # Đối với file docx, giả định là 1 trang gộp
                }
            )
            return [doc]
        except Exception as e:
            print(f"[Lỗi] Không thể nạp DOCX {file_path}: {str(e)}")
            return []
            
    def load_file(self, file_path: str, tag: str = None) -> List[Document]:
        """
        Tự động nhận diện định dạng file và gán nhãn tag (Metadata Tagging).
        """
        ext = os.path.splitext(file_path)[1].lower()
        docs = []
        
        if ext == '.pdf':
            docs = self.load_pdf(file_path)
        elif ext == '.docx':
            docs = self.load_docx(file_path)
        else:
            print(f"[Thông tin] Định dạng file {ext} chưa được hỗ trợ, bỏ qua: {file_path}")
            return []
            
        # Gán tag cho metadata để phục vụ tìm kiếm bộ lọc (Pre-filtering)
        if tag:
            for doc in docs:
                doc.metadata["doc_tag"] = tag
                
        return docs
        
    def load_directory(self, dir_path: str, tag: str = None) -> List[Document]:
        """
        Quét và nạp toàn bộ file PDF và Word trong thư mục chỉ định với tag tương ứng.
        """
        if not os.path.isdir(dir_path):
            print(f"[Cảnh báo] Thư mục không tồn tại: {dir_path}")
            return []
            
        all_docs = []
        # Tìm các file pdf và docx
        pdf_files = glob.glob(os.path.join(dir_path, "*.pdf"))
        docx_files = glob.glob(os.path.join(dir_path, "*.docx"))
        all_files = pdf_files + docx_files
        
        # Nếu không truyền tag, lấy tên thư mục con làm tag mặc định
        if not tag:
            tag = os.path.basename(os.path.normpath(dir_path))
            
        print(f"Bắt đầu nạp thư mục '{dir_path}' với nhãn tag='{tag}' ({len(all_files)} files)...")
        
        for file_path in all_files:
            docs = self.load_file(file_path, tag=tag)
            all_docs.extend(docs)
            
        return all_docs
        
    def load_all_folders(self, root_dir: str) -> List[Document]:
        """
        Quét thư mục gốc Ducument để tự động nạp các thư mục con và dán nhãn tương ứng.
        Ví dụ: Ducument/AI -> tag='AI', Ducument/LuatDatDai -> tag='LuatDatDai'.
        """
        if not os.path.exists(root_dir):
            print(f"[Cảnh báo] Thư mục gốc không tồn tại: {root_dir}")
            return []
            
        all_docs = []
        # Danh sách các thư mục con trực tiếp
        subdirs = [os.path.join(root_dir, d) for d in os.listdir(root_dir) 
                   if os.path.isdir(os.path.join(root_dir, d))]
                   
        for subdir in subdirs:
            folder_name = os.path.basename(subdir)
            docs = self.load_directory(subdir, tag=folder_name)
            all_docs.extend(docs)
            
        print(f"Hoàn thành nạp tổng cộng {len(all_docs)} trang tài liệu từ {root_dir}")
        return all_docs
